"""
build_corpus_mapping.py
=======================
构建 doc → excel 映射关系，并生成清洗后的训练语料。

功能:
1. 提取所有原始文档(PDF/Word)的文本
2. 提取所有Excel结构化数据
3. 建立文档段落 → Excel行的映射
4. 清洗原始语料(修复编码、去重、分句、标准化)
5. 输出: mapping.json + cleaned_corpus.jsonl

用法:
    python scripts/build_corpus_mapping.py
"""

import json
import os
import re
import sys
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import Optional

import openpyxl
import xlrd
import PyPDF2
from docx import Document

BASE = Path("/Users/kaitaocao/Library/CloudStorage/GoogleDrive-barrientosangie599@gmail.com"
            "/My Drive/AtlasNERDataset/国家电网文件")
OUTPUT_DIR = BASE / "清洗输出"

# ---------------------------------------------------------------------------
# 1. 文本提取
# ---------------------------------------------------------------------------

def extract_pdf_text(pdf_path: Path) -> list[dict]:
    """提取PDF每页文本，返回 [{page, text}]"""
    pages = []
    with open(pdf_path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                pages.append({"page": i + 1, "text": text})
    return pages


def extract_docx_text(docx_path: Path) -> list[dict]:
    """提取Word段落和表格文本"""
    doc = Document(str(docx_path))
    parts = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            parts.append({"type": "paragraph", "index": i, "text": text})

    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]
            row_text = " | ".join(c for c in cells if c)
            if row_text:
                parts.append({
                    "type": "table",
                    "table_index": ti,
                    "row_index": ri,
                    "text": row_text,
                    "cells": cells,
                })
    return parts


def extract_excel_rows(xlsx_path: Path) -> list[dict]:
    """提取Excel所有行（含表头）"""
    rows = []
    ext = xlsx_path.suffix.lower()

    if ext == ".xlsx":
        wb = openpyxl.load_workbook(str(xlsx_path), read_only=True)
        for sn in wb.sheetnames:
            ws = wb[sn]
            header = None
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                vals = [str(c).strip() if c is not None else "" for c in row]
                if i == 0 or (i == 1 and not header):
                    header = vals
                    continue
                if all(v == "" for v in vals):
                    continue
                record = {}
                for j, h in enumerate(header):
                    if j < len(vals) and h:
                        record[h] = vals[j]
                record["_sheet"] = sn
                record["_row"] = i + 1
                record["_source"] = xlsx_path.name
                rows.append(record)
        wb.close()

    elif ext == ".xls":
        wb = xlrd.open_workbook(str(xlsx_path))
        for sn in wb.sheet_names():
            ws = wb.sheet_by_name(sn)
            header = None
            for i in range(ws.nrows):
                vals = [str(ws.cell_value(i, j)).strip() for j in range(ws.ncols)]
                if i <= 1 and not header:
                    if any("序号" in v or "设备" in v or "操作" in v for v in vals):
                        header = vals
                        continue
                if header is None:
                    continue
                if all(v == "" or v == "0.0" for v in vals):
                    continue
                record = {}
                for j, h in enumerate(header):
                    if j < len(vals) and h:
                        record[h] = vals[j]
                record["_sheet"] = sn
                record["_row"] = i + 1
                record["_source"] = xlsx_path.name
                rows.append(record)
    return rows


# ---------------------------------------------------------------------------
# 2. 文本清洗
# ---------------------------------------------------------------------------

def fix_degree_symbol(text: str) -> str:
    """修复度数符号: 数字后跟。→ °（仅在角度上下文中）"""
    return re.sub(r'(\d+)。(～|~|-|—|到|至|\d)', r'\1°\2', text)


def normalize_whitespace(text: str) -> str:
    """标准化空白字符"""
    text = re.sub(r'[\r\n]+', '\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n ', '\n', text)
    return text.strip()


def remove_headers_footers(text: str) -> str:
    """去除PDF页眉页脚"""
    lines = text.split('\n')
    cleaned = []
    for line in lines:
        line_stripped = line.strip()
        if re.match(r'^[—\-]\s*\d+\s*[—\-]$', line_stripped):
            continue
        if re.match(r'^Q/GDW\s*\d+', line_stripped) and len(line_stripped) < 30:
            continue
        if re.match(r'^GB[/_]?T?\s*\d+', line_stripped) and len(line_stripped) < 30:
            continue
        if re.match(r'^\d+$', line_stripped) and len(line_stripped) <= 3:
            continue
        cleaned.append(line)
    return '\n'.join(cleaned)


def is_garbled(text: str) -> bool:
    """检测乱码文本 (如 /G44 /G77 编码损坏)"""
    garble_ratio = len(re.findall(r'/[A-Z]\d{2}', text)) / max(len(text), 1)
    if garble_ratio > 0.05:
        return True
    cjk_or_ascii = len(re.findall(r'[一-鿿　-〿a-zA-Z0-9]', text))
    if len(text) > 20 and cjk_or_ascii / len(text) < 0.3:
        return True
    return False


def split_sentences(text: str) -> list[str]:
    """中文分句"""
    text = normalize_whitespace(text)
    sents = re.split(r'([。！？；\n])', text)
    result = []
    current = ""
    for s in sents:
        current += s
        if s in ('。', '！', '？', '；', '\n'):
            sent = current.strip()
            if len(sent) >= 10 and not is_garbled(sent):
                if len(sent) > 512:
                    for chunk in [sent[i:i+512] for i in range(0, len(sent), 512)]:
                        if len(chunk) >= 10:
                            result.append(chunk)
                else:
                    result.append(sent)
            current = ""
    if current.strip() and len(current.strip()) >= 10 and not is_garbled(current.strip()):
        result.append(current.strip())
    return result


def clean_text(text: str) -> str:
    """综合清洗"""
    text = fix_degree_symbol(text)
    text = normalize_whitespace(text)
    text = remove_headers_footers(text)
    text = re.sub(r'　', ' ', text)  # 全角空格
    text = re.sub(r'　', ' ', text)
    text = re.sub(r'\s{2,}', ' ', text)
    return text.strip()


def clean_excel_cell(text: str) -> str:
    """清洗Excel单元格内容"""
    text = fix_degree_symbol(text)
    text = re.sub(r'[\r\n]+', '\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'　', ' ', text)
    return text.strip()


# ---------------------------------------------------------------------------
# 3. 映射构建
# ---------------------------------------------------------------------------

def build_phrase_index(excel_rows: list[dict]) -> dict[str, list[dict]]:
    """从Excel行中提取关键短语，建立倒排索引"""
    index = {}
    for row in excel_rows:
        for key, val in row.items():
            if key.startswith("_"):
                continue
            if not val or len(val) < 6:
                continue
            phrases = extract_key_phrases(val)
            for phrase in phrases:
                if phrase not in index:
                    index[phrase] = []
                index[phrase].append({
                    "source": row.get("_source", ""),
                    "row": row.get("_row", 0),
                    "sheet": row.get("_sheet", ""),
                    "field": key,
                })
    return index


def extract_key_phrases(text: str) -> list[str]:
    """提取用于匹配的关键短语（15-40字的片段）"""
    text = clean_excel_cell(text)
    phrases = []
    items = re.split(r'\d+[.、．]\s*', text)
    for item in items:
        item = item.strip()
        if len(item) >= 15:
            phrases.append(item[:40])
    if not phrases and len(text) >= 15:
        phrases.append(text[:40])
    return phrases


def find_matches(doc_text: str, phrase_index: dict, min_len: int = 15) -> list[dict]:
    """在文档文本中搜索Excel短语的匹配"""
    matches = []
    seen = set()
    for phrase, excel_refs in phrase_index.items():
        clean_phrase = phrase.replace(" ", "").replace("\n", "")
        clean_doc = doc_text.replace(" ", "").replace("\n", "")
        if len(clean_phrase) < min_len:
            continue
        if clean_phrase in clean_doc:
            key = (excel_refs[0]["source"], excel_refs[0]["row"], excel_refs[0]["field"])
            if key not in seen:
                seen.add(key)
                pos = clean_doc.index(clean_phrase)
                matches.append({
                    "phrase": phrase,
                    "doc_position": pos,
                    "excel_refs": excel_refs,
                })
    return matches


# ---------------------------------------------------------------------------
# 4. 主流程
# ---------------------------------------------------------------------------

def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # --- 收集所有原始文档 ---
    print("=" * 60)
    print("阶段1: 提取原始文档文本")
    print("=" * 60)

    all_docs = {}

    # PDF文件
    pdf_files = list(BASE.rglob("*.pdf"))
    for pdf_path in pdf_files:
        if "清洗输出" in str(pdf_path):
            continue
        rel = pdf_path.relative_to(BASE)
        print(f"  提取PDF: {rel}")
        pages = extract_pdf_text(pdf_path)
        full_text = "\n".join(p["text"] for p in pages)
        all_docs[str(rel)] = {
            "type": "pdf",
            "path": str(rel),
            "pages": len(pages),
            "total_chars": len(full_text),
            "text": full_text,
        }

    # Word文件
    docx_files = list(BASE.rglob("*.docx"))
    for docx_path in docx_files:
        if "清洗输出" in str(docx_path):
            continue
        rel = docx_path.relative_to(BASE)
        print(f"  提取Word: {rel}")
        parts = extract_docx_text(docx_path)
        full_text = "\n".join(p["text"] for p in parts)
        all_docs[str(rel)] = {
            "type": "docx",
            "path": str(rel),
            "parts": len(parts),
            "total_chars": len(full_text),
            "text": full_text,
        }

    print(f"\n  共提取 {len(all_docs)} 个文档")

    # --- 收集所有Excel数据 ---
    print("\n" + "=" * 60)
    print("阶段2: 提取Excel结构化数据")
    print("=" * 60)

    all_excel_rows = []
    excel_files = list(BASE.rglob("*.xlsx")) + list(BASE.rglob("*.xls"))
    for xlsx_path in excel_files:
        if "清洗输出" in str(xlsx_path):
            continue
        rel = xlsx_path.relative_to(BASE)
        print(f"  提取Excel: {rel}")
        rows = extract_excel_rows(xlsx_path)
        all_excel_rows.extend(rows)
        print(f"    → {len(rows)} 条记录")

    print(f"\n  共提取 {len(all_excel_rows)} 条Excel记录")

    # --- 清洗Excel数据 ---
    print("\n" + "=" * 60)
    print("阶段3: 清洗Excel数据")
    print("=" * 60)

    degree_fix_count = 0
    for row in all_excel_rows:
        for key, val in row.items():
            if key.startswith("_"):
                continue
            if not isinstance(val, str):
                continue
            cleaned = clean_excel_cell(val)
            if cleaned != val:
                if "°" in cleaned and "°" not in val:
                    degree_fix_count += 1
                row[key] = cleaned

    print(f"  修复度数符号: {degree_fix_count} 处")

    # --- 构建映射 ---
    print("\n" + "=" * 60)
    print("阶段4: 构建 doc → excel 映射")
    print("=" * 60)

    phrase_index = build_phrase_index(all_excel_rows)
    print(f"  短语索引: {len(phrase_index)} 条短语")

    mapping = {}
    for doc_key, doc_info in all_docs.items():
        doc_text = doc_info["text"]
        matches = find_matches(doc_text, phrase_index)
        if matches:
            excel_rows_matched = set()
            for m in matches:
                for ref in m["excel_refs"]:
                    excel_rows_matched.add((ref["source"], ref["row"]))

            mapping[doc_key] = {
                "doc_type": doc_info["type"],
                "doc_chars": doc_info["total_chars"],
                "match_count": len(matches),
                "unique_excel_rows_matched": len(excel_rows_matched),
                "matches": [
                    {
                        "phrase_preview": m["phrase"][:60],
                        "excel_source": m["excel_refs"][0]["source"],
                        "excel_row": m["excel_refs"][0]["row"],
                        "excel_field": m["excel_refs"][0]["field"],
                    }
                    for m in matches[:50]  # limit output size
                ],
            }
            print(f"  {doc_key}: {len(matches)} 匹配 → {len(excel_rows_matched)} 条Excel行")
        else:
            mapping[doc_key] = {
                "doc_type": doc_info["type"],
                "doc_chars": doc_info["total_chars"],
                "match_count": 0,
                "unique_excel_rows_matched": 0,
                "note": "无直接文本匹配（可能是间接引用/上游依据）",
            }
            print(f"  {doc_key}: 无直接匹配")

    # --- 生成清洗后语料 ---
    print("\n" + "=" * 60)
    print("阶段5: 生成清洗后训练语料")
    print("=" * 60)

    corpus_lines = []

    # 5a. Excel → 结构化句子
    for row in all_excel_rows:
        source = row.get("_source", "")
        row_num = row.get("_row", 0)

        text_fields = {}
        for key, val in row.items():
            if key.startswith("_") or not val or val in ("None", ""):
                continue
            text_fields[key] = val

        if not text_fields:
            continue

        # 把每行的各字段拼接为自然文本
        full_text = " ".join(f"{v}" for k, v in text_fields.items())
        full_text = clean_text(full_text)

        # 拆分为句子
        sentences = split_sentences(full_text)
        for sent in sentences:
            corpus_lines.append({
                "text": sent,
                "source_type": "excel",
                "source_file": source,
                "source_row": row_num,
                "fields": text_fields,
            })

    print(f"  Excel → {len(corpus_lines)} 句")

    # 5b. PDF/Word → 句子
    doc_sentence_count = 0
    for doc_key, doc_info in all_docs.items():
        text = doc_info["text"]
        text = clean_text(text)
        sentences = split_sentences(text)
        for sent in sentences:
            corpus_lines.append({
                "text": sent,
                "source_type": doc_info["type"],
                "source_file": doc_key,
            })
            doc_sentence_count += 1

    print(f"  PDF/Word → {doc_sentence_count} 句")
    print(f"  总计: {len(corpus_lines)} 句")

    # --- 去重 ---
    seen_texts = set()
    deduped = []
    for line in corpus_lines:
        normalized = re.sub(r'\s+', '', line["text"])
        if normalized not in seen_texts and len(normalized) >= 8:
            seen_texts.add(normalized)
            deduped.append(line)

    print(f"  去重后: {len(deduped)} 句 (去除 {len(corpus_lines) - len(deduped)} 重复)")

    # --- 输出 ---
    print("\n" + "=" * 60)
    print("阶段6: 写入输出文件")
    print("=" * 60)

    mapping_path = OUTPUT_DIR / "doc_excel_mapping.json"
    with open(mapping_path, "w", encoding="utf-8") as f:
        json.dump(mapping, f, ensure_ascii=False, indent=2)
    print(f"  映射文件: {mapping_path}")

    corpus_path = OUTPUT_DIR / "cleaned_corpus.jsonl"
    with open(corpus_path, "w", encoding="utf-8") as f:
        for line in deduped:
            f.write(json.dumps(line, ensure_ascii=False) + "\n")
    print(f"  清洗语料: {corpus_path}")

    excel_cleaned_path = OUTPUT_DIR / "cleaned_excel_rows.json"
    with open(excel_cleaned_path, "w", encoding="utf-8") as f:
        json.dump(all_excel_rows, f, ensure_ascii=False, indent=2)
    print(f"  清洗Excel: {excel_cleaned_path}")

    # --- 统计报告 ---
    print("\n" + "=" * 60)
    print("统计报告")
    print("=" * 60)

    total_matched_rows = 0
    for doc_key, info in mapping.items():
        total_matched_rows += info.get("unique_excel_rows_matched", 0)

    print(f"  原始文档数: {len(all_docs)}")
    print(f"  Excel总记录: {len(all_excel_rows)}")
    print(f"  有映射的文档: {sum(1 for v in mapping.values() if v['match_count'] > 0)}")
    print(f"  映射到的Excel行: {total_matched_rows}")
    print(f"  清洗后语料句数: {len(deduped)}")
    print(f"  输出目录: {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
