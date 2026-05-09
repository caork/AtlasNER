"""Locate source document pages for each Excel row.

For each labeled Excel row, finds the exact page(s) in matched source
documents that contain the row's content. Outputs an annotation manifest
that maps each row to its source page(s) with match quality scores.
"""

from __future__ import annotations

import json
import re
import unicodedata
from collections import defaultdict
from pathlib import Path

import fitz  # PyMuPDF
import docx

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------

GDRIVE = Path(
    "/Users/kaitaocao/Library/CloudStorage/"
    "GoogleDrive-barrientosangie599@gmail.com/My Drive/"
    "AtlasNERDataset/国家电网文件"
)
LABELED_ROWS = GDRIVE / "清洗输出" / "labeled_excel_rows.json"
OUTPUT_DIR = Path(__file__).resolve().parent.parent / "data" / "annotation"

DOC_SEARCH_DIRS = [
    GDRIVE / "模版及案例",
    GDRIVE / "补充来源文件",
]

# ---------------------------------------------------------------------------
# Document resolution
# ---------------------------------------------------------------------------


def resolve_doc_path(doc_name: str) -> Path | None:
    """Find the full path for a document name by searching known dirs."""
    for search_dir in DOC_SEARCH_DIRS:
        for p in search_dir.rglob(doc_name):
            return p
    return None


def build_doc_path_map(doc_names: set[str]) -> dict[str, Path]:
    """Resolve all document names to full paths."""
    result = {}
    for name in sorted(doc_names):
        path = resolve_doc_path(name)
        if path:
            result[name] = path
        else:
            print(f"  WARNING: could not find {name}")
    return result


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------


def normalize_text(text: str) -> str:
    """Normalize whitespace and unicode for matching."""
    text = unicodedata.normalize("NFKC", text)
    text = re.sub(r"\s+", "", text)
    return text


def extract_pdf_pages(pdf_path: Path) -> list[dict]:
    """Extract text from each page of a PDF."""
    pages = []
    doc = fitz.open(str(pdf_path))
    for i in range(doc.page_count):
        page = doc[i]
        text = page.get_text()
        pages.append({
            "page_num": i + 1,
            "raw_text": text,
            "normalized": normalize_text(text),
        })
    doc.close()
    return pages


def extract_docx_paragraphs(docx_path: Path) -> list[dict]:
    """Extract text from a DOCX, grouped into pseudo-pages by paragraph blocks.

    Since DOCX has no page concept without rendering, we group paragraphs
    into ~2000-char blocks as pseudo-pages.
    """
    doc = docx.Document(str(docx_path))
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    # Also extract table text
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                full_text += "\n" + " | ".join(row_texts)

    # Split into pseudo-pages of ~2000 chars at paragraph boundaries
    paragraphs = full_text.split("\n")
    pages = []
    current_text = ""
    page_num = 1

    for para in paragraphs:
        current_text += para + "\n"
        if len(current_text) >= 2000:
            pages.append({
                "page_num": page_num,
                "raw_text": current_text,
                "normalized": normalize_text(current_text),
            })
            current_text = ""
            page_num += 1

    if current_text.strip():
        pages.append({
            "page_num": page_num,
            "raw_text": current_text,
            "normalized": normalize_text(current_text),
        })

    return pages


def extract_pages(doc_path: Path) -> list[dict]:
    """Extract page-level text from any supported document."""
    suffix = doc_path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf_pages(doc_path)
    elif suffix == ".docx":
        return extract_docx_paragraphs(doc_path)
    else:
        print(f"  WARNING: unsupported format {suffix}: {doc_path.name}")
        return []


# ---------------------------------------------------------------------------
# Matching
# ---------------------------------------------------------------------------


def extract_match_phrases(row: dict) -> list[str]:
    """Extract key phrases from a row for matching against document pages.

    Uses prevention and control fields (longest, most unique text) as
    primary match signals, with equipment+process as secondary.
    """
    phrases = []

    # Primary: extract sub-phrases from prevention/control
    for field in ["风险防范措施", "工艺管控措施", "预控措施", "对应防范措施", "质量管控措施"]:
        val = row.get(field)
        if not val or str(val).strip().lower() == "nan":
            continue
        text = str(val).strip()
        # Extract individual numbered items as phrases
        items = re.split(r"\n\s*\d+[\.、．]", text)
        for item in items:
            item = item.strip()
            if len(item) >= 8:
                # Take first 30 chars as match phrase (avoid OCR noise at end)
                phrases.append(item[:min(50, len(item))])

    # Secondary: equipment + process as short identifiers
    equip = row.get("设备", "")
    proc = row.get("工序", "")
    if equip and str(equip).strip().lower() != "nan":
        phrases.append(str(equip).strip())
    if proc and str(proc).strip().lower() != "nan":
        phrases.append(str(proc).strip())

    return phrases


def score_page_match(page_normalized: str, phrases: list[str]) -> tuple[float, list[str]]:
    """Score how well a page matches the row's phrases.

    Returns (score, matched_phrases). Score is fraction of phrases found.
    """
    matched = []
    for phrase in phrases:
        norm_phrase = normalize_text(phrase)
        if len(norm_phrase) < 3:
            continue
        if norm_phrase in page_normalized:
            matched.append(phrase)
    score = len(matched) / max(len(phrases), 1)
    return score, matched


def find_best_pages(
    doc_pages: list[dict],
    phrases: list[str],
    min_score: float = 0.1,
) -> list[dict]:
    """Find the best matching pages in a document for given phrases."""
    results = []
    for page in doc_pages:
        score, matched = score_page_match(page["normalized"], phrases)
        if score >= min_score:
            results.append({
                "page_num": page["page_num"],
                "score": round(score, 3),
                "matched_phrases": len(matched),
                "total_phrases": len(phrases),
            })

    results.sort(key=lambda x: (-x["score"], x["page_num"]))
    return results


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------


def main() -> None:
    print(f"Loading labeled rows from {LABELED_ROWS}")
    with LABELED_ROWS.open("r", encoding="utf-8") as f:
        rows = json.load(f)
    print(f"  Total rows: {len(rows)}")

    # Collect all unique doc names
    all_doc_names: set[str] = set()
    for row in rows:
        for d in row.get("_matched_docs", []):
            all_doc_names.add(d)
    print(f"  Unique documents: {len(all_doc_names)}")

    # Resolve paths
    doc_paths = build_doc_path_map(all_doc_names)
    print(f"  Resolved: {len(doc_paths)}/{len(all_doc_names)}")

    # Extract page text for all documents (cached)
    print("\nExtracting document pages...")
    doc_pages_cache: dict[str, list[dict]] = {}
    for doc_name, doc_path in sorted(doc_paths.items()):
        print(f"  {doc_name} ...", end=" ", flush=True)
        pages = extract_pages(doc_path)
        doc_pages_cache[doc_name] = pages
        print(f"{len(pages)} pages")

    # Match each row to pages
    print("\nLocating passages...")
    manifest_entries: list[dict] = []
    stats = defaultdict(int)

    for idx, row in enumerate(rows):
        matched_docs = row.get("_matched_docs", [])
        entry = {
            "row_index": idx,
            "row_num": row.get("_row", 0),
            "source_file": row.get("_source", "unknown"),
            "source_sheet": row.get("_sheet", "unknown"),
            "tier": row.get("_tier", "unknown"),
            "equipment": row.get("设备", ""),
            "process": row.get("工序", ""),
            "doc_matches": [],
            "status": "unmatched",
        }

        if not matched_docs:
            stats["unmatched"] += 1
            entry["status"] = "unmatched"
            manifest_entries.append(entry)
            continue

        phrases = extract_match_phrases(row)
        if not phrases:
            stats["no_phrases"] += 1
            entry["status"] = "no_phrases"
            manifest_entries.append(entry)
            continue

        best_overall_score = 0.0
        for doc_name in matched_docs:
            if doc_name not in doc_pages_cache:
                continue
            pages = doc_pages_cache[doc_name]
            page_matches = find_best_pages(pages, phrases)
            if page_matches:
                best_score = page_matches[0]["score"]
                best_overall_score = max(best_overall_score, best_score)
                entry["doc_matches"].append({
                    "doc_name": doc_name,
                    "doc_type": Path(doc_name).suffix.lower(),
                    "best_score": best_score,
                    "pages": page_matches[:5],
                })

        if entry["doc_matches"]:
            entry["doc_matches"].sort(key=lambda x: -x["best_score"])
            entry["status"] = "matched"
            if best_overall_score >= 0.5:
                stats["high_match"] += 1
            elif best_overall_score >= 0.2:
                stats["medium_match"] += 1
            else:
                stats["low_match"] += 1
        else:
            stats["match_failed"] += 1
            entry["status"] = "match_failed"

        manifest_entries.append(entry)

        if (idx + 1) % 200 == 0:
            print(f"  Processed {idx + 1}/{len(rows)}...")

    # Write manifest
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    manifest_path = OUTPUT_DIR / "annotation_manifest.json"
    manifest = {
        "total_rows": len(rows),
        "stats": dict(stats),
        "entries": manifest_entries,
    }
    with manifest_path.open("w", encoding="utf-8") as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)

    print(f"\n=== Results ===")
    print(f"  Total: {len(rows)}")
    for k, v in sorted(stats.items()):
        print(f"  {k}: {v}")
    print(f"\n  Wrote {manifest_path}")


if __name__ == "__main__":
    main()
